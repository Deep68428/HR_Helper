from pathlib import Path

import pdfplumber
from docx import Document
from loguru import logger

from backend.core.deps import get_ocr

ocr = get_ocr()
# ================= TEXT EXTRACTION =================


def extract_text_from_docx(path: str) -> str:
    logger.info(f"📄 Extracting text from DOCX: {Path(path).name}")

    try:
        doc = Document(path)
        chunks = []

        # ---- HEADERS ----
        for section in doc.sections:
            header = section.header
            for p in header.paragraphs:
                if p.text.strip():
                    chunks.append(p.text.strip())

            footer = section.footer
            for p in footer.paragraphs:
                if p.text.strip():
                    chunks.append(p.text.strip())

        # ---- BODY PARAGRAPHS ----
        for p in doc.paragraphs:
            if p.text.strip():
                chunks.append(p.text.strip())

        # ---- TABLES ----
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    if c.text.strip():
                        chunks.append(c.text.strip())

        text = "\n".join(dict.fromkeys(chunks))
        logger.info(f"✅ Successfully extracted {len(text)} characters from DOCX")
        return text

    except Exception as e:
        logger.error(f"❌ Failed to extract text from DOCX {Path(path).name}: {e}")
        raise


def extract_text_from_pdf(path: str) -> str:
    logger.info(f"📄 Extracting text from PDF: {Path(path).name}")

    try:
        text = ""
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            logger.debug(f"Processing {page_count} pages with pdfplumber")

            for i, page in enumerate(pdf.pages, 1):
                t = page.extract_text()
                if t:
                    text += t + "\n"
                    logger.debug(
                        f"Page {i}/{page_count}: extracted {len(t)} characters"
                    )

        # OCR fallback if pdfplumber didn't extract enough text
        if len(text.strip()) < 50:
            logger.warning(
                f"⚠️  Insufficient text from pdfplumber ({len(text.strip())} chars), falling back to OCR"
            )
            result = ocr.predict(path)
            lines = []
            for page in result:
                for box in page:
                    lines.append(box[1][0])
            text = "\n".join(lines)
            logger.info(
                f"🔍 OCR extracted {len(text)} characters from {len(result)} pages"
            )

        final_text = text.strip()
        logger.info(f"✅ Successfully extracted {len(final_text)} characters from PDF")
        return final_text

    except Exception as e:
        logger.error(f"❌ Failed to extract text from PDF {Path(path).name}: {e}")
        raise


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    logger.info(f"🔍 Starting text extraction for file type: {ext}")

    if ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".pdf":
        return extract_text_from_pdf(path)
    else:
        logger.error(f"❌ Unsupported file type: {ext}")
        raise ValueError(f"Unsupported file type: {ext}")
